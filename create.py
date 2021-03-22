import re
from docx import Document
import os
from urllib.parse import urlparse


def mails(csv, templ):
    all_mails = []
    data = csv['data']

    for row in data:
        m = {}

        # Subject
        m['subject'] = row['oggetto'] 

        # Recipient
        m['recipient'] = row['pec']

        # Body
        doc = mail_text(row, templ)
        text_doc = get_text(doc)
        m['body'] = text_doc

        # Attachments
        m['attachments'] = []
        for k,v in row.items():
            if 'allegato' in k and v != '':
                m['attachments'].append(v)

        all_mails.append(m)
    
    return all_mails


def mail_text(row, templ):
    doc = Document()
    pattern = '\$\{(.*?)\}'

    for p in templ.paragraphs:
        par_vars = re.findall(pattern, p.text)

        # Empty line
        if p.text == '':
            continue
        
        # No variables
        if len(par_vars) == 0:
            doc.add_paragraph(p.text)
            continue
        
        # Replace vars
        upd_p = p.text
        for v in par_vars:
            upd_p = upd_p.replace('${' + v + '}', row[v])
        
        doc.add_paragraph(upd_p)
        
    return doc


def get_text(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    return '\n\n'.join(fullText)
